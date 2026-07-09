from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def normalize_value(val):
    if isinstance(val, list):
        return ", ".join(str(x) for x in val)
    if isinstance(val, tuple):
        return ", ".join(str(x) for x in val)
    if isinstance(val, set):
        return ", ".join(str(x) for x in val)
    if isinstance(val, dict):
        return str(val)
    if val is None:
        return ""
    return str(val)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def build_word(conversation, database: str = "") -> BytesIO:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI Database Agent")
    r.bold = True
    r.font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Conversație completă exportată")
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Bază de date: {database or '-'}")
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Export: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    r.font.size = Pt(11)

    doc.add_paragraph("")

    for idx, item in enumerate(conversation, start=1):
        classification = item.get("classification", {}) or {}
        question = normalize_value(item.get("question", ""))
        tip = normalize_value(classification.get("type", classification.get("tip", "")))
        kind = normalize_value(classification.get("kind", ""))
        intent = normalize_value(classification.get("intent", ""))
        sql_text = normalize_value(item.get("sql", ""))
        explanation = normalize_value(item.get("answer", ""))
        result = item.get("result")

        doc.add_heading(f"Interacțiunea {idx}", level=1)

        p = doc.add_paragraph()
        p.add_run("Întrebare: ").bold = True
        p.add_run(question or "-")

        p = doc.add_paragraph()
        p.add_run("Tip: ").bold = True
        p.add_run(tip or "-")

        p = doc.add_paragraph()
        p.add_run("Kind: ").bold = True
        p.add_run(kind or "-")

        p = doc.add_paragraph()
        p.add_run("Intent: ").bold = True
        p.add_run(intent or "-")

        if sql_text:
            p = doc.add_paragraph()
            p.add_run("SQL:").bold = True
            doc.add_paragraph(sql_text)

        if explanation:
            p = doc.add_paragraph()
            p.add_run("Explicație:").bold = True
            doc.add_paragraph(explanation)

        if isinstance(result, dict) and result.get("columns") and result.get("rows") is not None:
            columns = [normalize_value(c) for c in result.get("columns", [])]
            rows = result.get("rows", [])[:10]

            doc.add_paragraph("Preview rezultate:")

            landscape_section = doc.add_section(WD_SECTION.NEW_PAGE)
            landscape_section.orientation = WD_ORIENT.LANDSCAPE
            landscape_section.page_width, landscape_section.page_height = (
                landscape_section.page_height,
                landscape_section.page_width,
            )
            landscape_section.top_margin = Inches(0.5)
            landscape_section.bottom_margin = Inches(0.5)
            landscape_section.left_margin = Inches(0.5)
            landscape_section.right_margin = Inches(0.5)

            table = doc.add_table(rows=1, cols=len(columns))
            table.style = "Table Grid"

            header_cells = table.rows[0].cells
            for i, col in enumerate(columns):
                header_cells[i].text = col
                shade_cell(header_cells[i], "2F5597")
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            set_repeat_table_header(table.rows[0])

            for row_idx, row in enumerate(rows, start=1):
                if isinstance(row, dict):
                    values = [normalize_value(row.get(col, "")) for col in columns]
                else:
                    values = [normalize_value(v) for v in row]

                cells = table.add_row().cells
                for i, value in enumerate(values[:len(columns)]):
                    text = value
                    if len(text) > 60:
                        text = text[:57] + "..."
                    cells[i].text = text
                    if row_idx % 2:
                        shade_cell(cells[i], "F8FBFF")
                    else:
                        shade_cell(cells[i], "EDF3FA")

            p = doc.add_paragraph()
            p.add_run(
                f"Rânduri returnate total: {result.get('total_rows', len(result.get('rows', [])))} | "
                f"În document sunt afișate primele {len(rows)}."
            )

            portrait_section = doc.add_section(WD_SECTION.NEW_PAGE)
            portrait_section.orientation = WD_ORIENT.PORTRAIT
            portrait_section.page_width, portrait_section.page_height = (
                section.page_width,
                section.page_height,
            )
            portrait_section.top_margin = Inches(0.7)
            portrait_section.bottom_margin = Inches(0.7)
            portrait_section.left_margin = Inches(0.75)
            portrait_section.right_margin = Inches(0.75)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio