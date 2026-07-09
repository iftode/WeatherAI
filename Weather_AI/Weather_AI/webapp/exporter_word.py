from io import BytesIO
from datetime import datetime
from docx import Document


def build_word(conversation, database: str = "") -> BytesIO:
    doc = Document()

    doc.add_heading("AI Database Agent — Conversație completă", level=1)
    doc.add_paragraph(f"Database: {database or ''}")
    doc.add_paragraph(f"Exported: {datetime.now().isoformat(timespec='seconds')}")
    doc.add_paragraph("")

    for idx, item in enumerate(conversation, start=1):
        doc.add_heading(f"Interacțiunea {idx}", level=2)

        doc.add_paragraph("Întrebare utilizator:")
        doc.add_paragraph(item.get("question", ""))

        classification = item.get("classification", {})
        kind = classification.get("kind", "")
        intent = classification.get("intent", "")
        doc.add_paragraph(f"Clasificare: {kind} | Intent: {intent}")

        if item.get("type") == "business":
            if item.get("complexity"):
                doc.add_paragraph(f"Complexitate răspuns: {item.get('complexity')}")
            doc.add_paragraph("Răspuns:")
            doc.add_paragraph(item.get("answer", ""))

        elif item.get("type") == "data":
            doc.add_paragraph("SQL generat:")
            doc.add_paragraph(item.get("sql", ""))

            if item.get("explanation"):
                doc.add_paragraph("Explicație SQL:")
                doc.add_paragraph(item.get("explanation", ""))

            columns = item.get("columns", [])
            rows = item.get("rows", [])

            doc.add_paragraph("Rezultate:")

            if not columns:
                doc.add_paragraph("(Fără rezultate)")
            else:
                table = doc.add_table(rows=1, cols=len(columns))
                table.style = "Table Grid"

                hdr = table.rows[0].cells
                for i, col in enumerate(columns):
                    hdr[i].text = str(col)

                for r in rows:
                    cells = table.add_row().cells
                    for i, val in enumerate(r):
                        cells[i].text = "" if val is None else str(val)

                doc.add_paragraph(f"Rânduri returnate: {len(rows)}")

        doc.add_paragraph("-" * 60)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio